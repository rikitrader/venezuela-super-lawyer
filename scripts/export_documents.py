#!/usr/bin/env python3
"""
Venezuela Super Lawyer - Document Export Module
Exports legal reports and documents to PDF and DOCX formats.

Version: 1.0.0
"""

__version__ = "1.0.0"
__author__ = "Venezuela Super Lawyer"

import os
import sys
import re
import argparse
from enum import Enum
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any


class ExportFormat(Enum):
    """Supported export formats."""
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"


# Check for optional dependencies
try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    from weasyprint import HTML as WeasyHTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ═══════════════════════════════════════════════════════════════════════════════
#                         CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════

# Output directories
REPORTS_DIR = Path(__file__).parent.parent / "reportes_legales"
EXPORTS_DIR = Path(__file__).parent.parent / "exports"

# PDF styling
PDF_CSS = """
@page {
    size: letter;
    margin: 2.5cm;
    @top-center {
        content: "VENEZUELA SUPER LAWYER";
        font-size: 10pt;
        color: #666;
    }
    @bottom-center {
        content: "Página " counter(page) " de " counter(pages);
        font-size: 10pt;
        color: #666;
    }
}

body {
    font-family: 'Times New Roman', Times, serif;
    font-size: 12pt;
    line-height: 1.6;
    color: #333;
}

h1 {
    font-size: 18pt;
    color: #1a1a1a;
    border-bottom: 2px solid #003366;
    padding-bottom: 10px;
    margin-top: 20px;
}

h2 {
    font-size: 14pt;
    color: #003366;
    margin-top: 15px;
}

h3 {
    font-size: 12pt;
    color: #003366;
    font-style: italic;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 15px 0;
}

th, td {
    border: 1px solid #ccc;
    padding: 8px;
    text-align: left;
}

th {
    background-color: #003366;
    color: white;
}

tr:nth-child(even) {
    background-color: #f9f9f9;
}

code {
    font-family: 'Courier New', monospace;
    background-color: #f4f4f4;
    padding: 2px 5px;
    border-radius: 3px;
}

pre {
    background-color: #f4f4f4;
    padding: 15px;
    border-radius: 5px;
    overflow-x: auto;
    font-size: 10pt;
}

blockquote {
    border-left: 4px solid #003366;
    margin: 15px 0;
    padding: 10px 20px;
    background-color: #f9f9f9;
    font-style: italic;
}

.warning {
    background-color: #fff3cd;
    border: 1px solid #ffc107;
    padding: 10px;
    border-radius: 5px;
    margin: 10px 0;
}

.critical {
    background-color: #f8d7da;
    border: 1px solid #dc3545;
    padding: 10px;
    border-radius: 5px;
    margin: 10px 0;
}

.success {
    background-color: #d4edda;
    border: 1px solid #28a745;
    padding: 10px;
    border-radius: 5px;
    margin: 10px 0;
}

.header-banner {
    background-color: #003366;
    color: white;
    padding: 20px;
    text-align: center;
    margin-bottom: 20px;
}

.footer {
    margin-top: 30px;
    padding-top: 10px;
    border-top: 1px solid #ccc;
    font-size: 10pt;
    color: #666;
}
"""


def get_pdf_css() -> str:
    """Return the CSS styles for PDF generation."""
    return PDF_CSS


def get_docx_styles() -> Dict[str, Any]:
    """Return style configuration for DOCX generation."""
    return {
        "heading1": {"font": "Times New Roman", "size": 18, "bold": True},
        "heading2": {"font": "Times New Roman", "size": 14, "bold": True},
        "heading3": {"font": "Times New Roman", "size": 12, "bold": True, "italic": True},
        "normal": {"font": "Times New Roman", "size": 12},
        "title": {"font": "Times New Roman", "size": 24, "bold": True},
        "footer": {"font": "Times New Roman", "size": 10, "color": "#666666"}
    }


# ═══════════════════════════════════════════════════════════════════════════════
#                         MARKDOWN TO HTML
# ═══════════════════════════════════════════════════════════════════════════════

def markdown_to_html(markdown_text: str) -> str:
    """
    Convert Markdown to HTML.
    Uses markdown library if available, otherwise basic conversion.
    """
    try:
        import markdown
        html = markdown.markdown(
            markdown_text,
            extensions=[
                'markdown.extensions.tables',
                'markdown.extensions.fenced_code',
                'markdown.extensions.toc',
                'markdown.extensions.meta',
            ]
        )
        return html
    except ImportError:
        # Fallback: basic markdown conversion
        return _basic_markdown_to_html(markdown_text)


def _basic_markdown_to_html(text: str) -> str:
    """Basic Markdown to HTML conversion without external libraries."""
    html = text

    # Headers
    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)

    # Bold and italic
    html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
    html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)

    # Code blocks
    html = re.sub(r'```(\w*)\n(.*?)```', r'<pre><code>\2</code></pre>', html, flags=re.DOTALL)
    html = re.sub(r'`([^`]+)`', r'<code>\1</code>', html)

    # Lists
    html = re.sub(r'^- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
    html = re.sub(r'(<li>.*</li>\n)+', r'<ul>\g<0></ul>', html)

    # Tables (basic support)
    lines = html.split('\n')
    in_table = False
    table_lines = []
    result_lines = []

    for line in lines:
        if '|' in line and not line.strip().startswith('```'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
        else:
            if in_table:
                result_lines.append(_convert_table(table_lines))
                in_table = False
                table_lines = []
            result_lines.append(line)

    if in_table:
        result_lines.append(_convert_table(table_lines))

    html = '\n'.join(result_lines)

    # Paragraphs
    html = re.sub(r'\n\n+', '</p><p>', html)
    html = f'<p>{html}</p>'

    # Clean up
    html = re.sub(r'<p>\s*</p>', '', html)
    html = re.sub(r'<p>\s*<(h[1-6]|ul|ol|table|pre)', r'<\1', html)
    html = re.sub(r'</(h[1-6]|ul|ol|table|pre)>\s*</p>', r'</\1>', html)

    return html


def _convert_table(lines: list) -> str:
    """Convert markdown table lines to HTML table."""
    if len(lines) < 2:
        return '\n'.join(lines)

    html = '<table>\n'

    # Header
    header_cells = [c.strip() for c in lines[0].split('|') if c.strip()]
    html += '<thead><tr>'
    for cell in header_cells:
        html += f'<th>{cell}</th>'
    html += '</tr></thead>\n'

    # Body (skip separator line)
    html += '<tbody>'
    for line in lines[2:]:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        html += '<tr>'
        for cell in cells:
            html += f'<td>{cell}</td>'
        html += '</tr>\n'
    html += '</tbody></table>'

    return html


# ═══════════════════════════════════════════════════════════════════════════════
#                         PDF EXPORT
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_pdf(
    markdown_file: str,
    output_file: Optional[str] = None,
    title: Optional[str] = None
) -> Optional[str]:
    """
    Export a Markdown file to PDF.

    Args:
        markdown_file: Path to the markdown file
        output_file: Optional output PDF path (default: same name with .pdf)
        title: Optional document title

    Returns:
        Path to the generated PDF, or None if failed
    """
    markdown_path = Path(markdown_file)

    if not markdown_path.exists():
        print(f"Error: File not found: {markdown_file}", file=sys.stderr)
        return None

    # Read markdown content
    content = markdown_path.read_text(encoding='utf-8')

    # Generate title from filename if not provided
    if not title:
        title = markdown_path.stem.replace('_', ' ').title()

    # Convert to HTML
    html_body = markdown_to_html(content)

    # Build full HTML document
    html_document = f"""
    <!DOCTYPE html>
    <html lang="es">
    <head>
        <meta charset="UTF-8">
        <title>{title}</title>
        <style>
            {PDF_CSS}
        </style>
    </head>
    <body>
        <div class="header-banner">
            <h1 style="color: white; border: none; margin: 0;">{title}</h1>
            <p style="margin: 5px 0;">Venezuela Super Lawyer - Sistema Legal AI</p>
            <p style="margin: 0; font-size: 10pt;">Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
        </div>
        {html_body}
        <div class="footer">
            <p>Documento generado por Venezuela Super Lawyer v1.0.0</p>
            <p>Este documento es para fines informativos y no constituye asesoría legal.</p>
        </div>
    </body>
    </html>
    """

    # Determine output path
    if output_file:
        output_path = Path(output_file)
    else:
        EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
        output_path = EXPORTS_DIR / f"{markdown_path.stem}.pdf"

    # Try weasyprint first
    try:
        from weasyprint import HTML, CSS
        HTML(string=html_document).write_pdf(str(output_path))
        print(f"✓ PDF exported to: {output_path}")
        return str(output_path)
    except ImportError:
        pass

    # Fallback: save HTML (can be opened in browser and printed to PDF)
    html_output = output_path.with_suffix('.html')
    html_output.write_text(html_document, encoding='utf-8')
    print(f"⚠ weasyprint not installed. HTML saved to: {html_output}")
    print("  Install with: pip install weasyprint")
    print("  Or open the HTML file in a browser and print to PDF.")
    return str(html_output)


# ═══════════════════════════════════════════════════════════════════════════════
#                         DOCX EXPORT
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_docx(
    markdown_file: str,
    output_file: Optional[str] = None,
    title: Optional[str] = None
) -> Optional[str]:
    """
    Export a Markdown file to DOCX (Microsoft Word).

    Args:
        markdown_file: Path to the markdown file
        output_file: Optional output DOCX path (default: same name with .docx)
        title: Optional document title

    Returns:
        Path to the generated DOCX, or None if failed
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("Error: python-docx not installed.", file=sys.stderr)
        print("Install with: pip install python-docx", file=sys.stderr)
        return None

    markdown_path = Path(markdown_file)

    if not markdown_path.exists():
        print(f"Error: File not found: {markdown_file}", file=sys.stderr)
        return None

    # Read markdown content
    content = markdown_path.read_text(encoding='utf-8')

    # Generate title from filename if not provided
    if not title:
        title = markdown_path.stem.replace('_', ' ').title()

    # Create document
    doc = Document()

    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Venezuela Super Lawyer - Sistema Legal AI\n").italic = True
    meta.add_run(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    doc.add_paragraph()  # Spacer

    # Parse and add content
    lines = content.split('\n')
    in_code_block = False
    code_lines = []

    for line in lines:
        # Skip ASCII art headers
        if '═' in line or '╔' in line or '╚' in line or '║' in line:
            continue

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                para = doc.add_paragraph()
                run = para.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                code_lines = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Headers
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        # List items
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
        # Tables (simplified - add as paragraphs)
        elif '|' in line and not line.startswith('|--'):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                doc.add_paragraph(' | '.join(cells))
        # Regular paragraphs
        elif line.strip():
            para = doc.add_paragraph()
            # Handle bold and italic
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                else:
                    para.add_run(part)

    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('─' * 50)
    footer = doc.add_paragraph()
    footer.add_run("Documento generado por Venezuela Super Lawyer v1.0.0\n").italic = True
    footer.add_run("Este documento es para fines informativos y no constituye asesoría legal.").italic = True

    # Determine output path
    if output_file:
        output_path = Path(output_file)
    else:
        EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
        output_path = EXPORTS_DIR / f"{markdown_path.stem}.docx"

    # Save document
    doc.save(str(output_path))
    print(f"✓ DOCX exported to: {output_path}")
    return str(output_path)


# ═══════════════════════════════════════════════════════════════════════════════
#                         BATCH EXPORT
# ═══════════════════════════════════════════════════════════════════════════════

def export_all_reports(format: str = "pdf") -> list:
    """
    Export all reports in reportes_legales directory.

    Args:
        format: "pdf", "docx", or "both"

    Returns:
        List of exported file paths
    """
    exported = []

    if not REPORTS_DIR.exists():
        print(f"Reports directory not found: {REPORTS_DIR}")
        return exported

    markdown_files = list(REPORTS_DIR.glob("*.md"))

    if not markdown_files:
        print("No markdown files found in reports directory.")
        return exported

    print(f"Found {len(markdown_files)} report(s) to export...")

    for md_file in markdown_files:
        if md_file.name in ('INDEX.md', '.gitkeep'):
            continue

        print(f"\nExporting: {md_file.name}")

        if format in ("pdf", "both"):
            result = export_to_pdf(str(md_file))
            if result:
                exported.append(result)

        if format in ("docx", "both"):
            result = export_to_docx(str(md_file))
            if result:
                exported.append(result)

    print(f"\n✓ Exported {len(exported)} file(s)")
    return exported


# ═══════════════════════════════════════════════════════════════════════════════
#                         CLI INTERFACE
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    """CLI interface for document export."""
    import argparse

    parser = argparse.ArgumentParser(
        description="Venezuela Super Lawyer - Document Export",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python3 export_documents.py pdf report.md
  python3 export_documents.py docx report.md -o output.docx
  python3 export_documents.py pdf report.md --title "Informe Legal"
  python3 export_documents.py all --format both
  python3 export_documents.py all --format pdf
        """
    )

    subparsers = parser.add_subparsers(dest="command", help="Export commands")

    # PDF export
    pdf_parser = subparsers.add_parser("pdf", help="Export to PDF")
    pdf_parser.add_argument("file", help="Markdown file to export")
    pdf_parser.add_argument("-o", "--output", help="Output file path")
    pdf_parser.add_argument("-t", "--title", help="Document title")

    # DOCX export
    docx_parser = subparsers.add_parser("docx", help="Export to DOCX")
    docx_parser.add_argument("file", help="Markdown file to export")
    docx_parser.add_argument("-o", "--output", help="Output file path")
    docx_parser.add_argument("-t", "--title", help="Document title")

    # Batch export all reports
    all_parser = subparsers.add_parser("all", help="Export all reports")
    all_parser.add_argument(
        "--format",
        choices=["pdf", "docx", "both"],
        default="pdf",
        help="Export format (default: pdf)"
    )

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        return

    if args.command == "pdf":
        export_to_pdf(args.file, args.output, args.title)

    elif args.command == "docx":
        export_to_docx(args.file, args.output, args.title)

    elif args.command == "all":
        export_all_reports(args.format)


if __name__ == "__main__":
    main()
